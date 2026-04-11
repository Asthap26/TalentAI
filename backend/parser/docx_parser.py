import docx

def extract_text_from_docx(filepath: str) -> str:
    """
    Extracts text from a DOCX file, including text from standard paragraphs
    and embedded tables.
    """
    text = ""
    try:
        doc = docx.Document(filepath)
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
                
        # Extract text from tables which are often used for layout in resumes
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
    return text

if __name__ == "__main__":
    print("DOCX Parser initialized.")
